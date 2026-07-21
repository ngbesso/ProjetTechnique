import io

import httpx
from docx import Document
from openpyxl import Workbook
from pydantic import BaseModel
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from app.core.config import settings

DOMAIN_LABELS = {
    "membres": "Membres",
    "dons": "Dons",
    "evenements": "Événements et formations",
    "sermons": "Sermons",
    "articles": "Blog",
    "eglises": "Églises affiliées",
}


def flatten_stats(
    stats: BaseModel,
) -> tuple[list[tuple[str, str]], dict[str, list[dict]]]:
    """Aplatit un objet de statistiques Pydantic en un résumé clé/valeur (champs
    scalaires) et des tables nommées (champs liste), quel que soit le domaine."""
    summary: list[tuple[str, str]] = []
    tables: dict[str, list[dict]] = {}
    for name, value in stats.model_dump().items():
        if isinstance(value, list):
            tables[name] = value
        else:
            summary.append((name, str(value)))
    return summary, tables


def fetch_ai_summary(domain_label: str, authorization: str) -> str | None:
    """Demande à l'assistant IA admin une courte synthèse du domaine. Meilleur
    effort : retourne None sur toute erreur ou dépassement de délai plutôt que de
    faire échouer la génération du rapport."""
    question = (
        f"Rédige une synthèse en 3 phrases maximum des statistiques du domaine "
        f"« {domain_label} »."
    )
    try:
        res = httpx.post(
            f"{settings.ai_service_url}/admin/chat",
            json={"question": question},
            headers={"Authorization": authorization},
            timeout=30.0,
        )
        res.raise_for_status()
        return res.json()["answer"]
    except (httpx.HTTPError, KeyError, ValueError):
        return None


def build_excel(
    title: str,
    summary: list[tuple[str, str]],
    tables: dict[str, list[dict]],
    ai_summary: str | None = None,
) -> bytes:
    wb = Workbook()
    ws = wb.active
    ws.title = "Résumé"
    ws.append([title])
    ws.append([])
    if ai_summary:
        ws.append(["Synthèse IA", ai_summary])
        ws.append([])
    for key, value in summary:
        ws.append([key, value])

    for name, rows in tables.items():
        sheet = wb.create_sheet(title=name[:31])
        if rows:
            headers = list(rows[0].keys())
            sheet.append(headers)
            for row in rows:
                sheet.append([row.get(h) for h in headers])

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def build_word(
    title: str,
    summary: list[tuple[str, str]],
    tables: dict[str, list[dict]],
    ai_summary: str | None = None,
) -> bytes:
    doc = Document()
    doc.add_heading(title, level=1)

    if ai_summary:
        doc.add_heading("Synthèse IA", level=2)
        doc.add_paragraph(ai_summary)

    doc.add_heading("Résumé", level=2)
    summary_table = doc.add_table(rows=0, cols=2)
    summary_table.style = "Light Grid Accent 1"
    for key, value in summary:
        row = summary_table.add_row().cells
        row[0].text = key
        row[1].text = value

    for name, rows in tables.items():
        doc.add_heading(name, level=2)
        if not rows:
            doc.add_paragraph("Aucune donnée.")
            continue
        headers = list(rows[0].keys())
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Light Grid Accent 1"
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
        for row in rows:
            cells = table.add_row().cells
            for i, h in enumerate(headers):
                cells[i].text = str(row.get(h, ""))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_pdf(
    title: str,
    summary: list[tuple[str, str]],
    tables: dict[str, list[dict]],
    ai_summary: str | None = None,
) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter)
    styles = getSampleStyleSheet()
    story = [Paragraph(title, styles["Title"]), Spacer(1, 12)]

    if ai_summary:
        story.append(Paragraph("Synthèse IA", styles["Heading2"]))
        story.append(Paragraph(ai_summary, styles["Normal"]))
        story.append(Spacer(1, 12))

    story.append(Paragraph("Résumé", styles["Heading2"]))
    summary_data = [[key, value] for key, value in summary]
    if summary_data:
        story.append(_styled_table(summary_data))
    story.append(Spacer(1, 12))

    for name, rows in tables.items():
        story.append(Paragraph(name, styles["Heading2"]))
        if not rows:
            story.append(Paragraph("Aucune donnée.", styles["Normal"]))
        else:
            headers = list(rows[0].keys())
            data = [headers] + [[str(row.get(h, "")) for h in headers] for row in rows]
            story.append(_styled_table(data))
        story.append(Spacer(1, 12))

    doc.build(story)
    return buf.getvalue()


def _styled_table(data: list[list[str]]) -> Table:
    table = Table(data)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#3b2f8a")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ]
        )
    )
    return table
